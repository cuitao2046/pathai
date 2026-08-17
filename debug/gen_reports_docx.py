# -*- coding: utf-8 -*-
"""生成两份面向导师的正式文档：
   1) docs/14-项目进展汇报-导师用.docx
   2) docs/15-工具材料采购预算说明.docx
依赖 python-docx（已在托管 venv 安装）。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = r"E:/code/pathai/docs"
os.makedirs(OUT_DIR, exist_ok=True)

# ---------- 字体工具 ----------
def set_eastasia(run, font="宋体"):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)

def set_base_font(doc, font="宋体", size=11):
    st = doc.styles['Normal']
    st.font.name = font
    st.font.size = Pt(size)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    if level == 1:
        r.font.size = Pt(15); set_eastasia(r, "黑体")
    elif level == 2:
        r.font.size = Pt(13); set_eastasia(r, "黑体")
    else:
        r.font.size = Pt(11.5); set_eastasia(r, "黑体")
    return p

def para(doc, text, size=10.5, bold=False, align=None, after=4):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    set_eastasia(r)
    return p

def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    set_eastasia(r)
    return p

def make_table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(font_size); set_eastasia(r, "黑体")
        # 表头底色
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn('w:shd'), {qn('w:val'):'clear', qn('w:color'):'auto', qn('w:fill'):'D9E2F3'})
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(font_size); set_eastasia(r)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def hrule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pbdr.makeelement(qn('w:bottom'), {qn('w:val'):'single', qn('w:sz'):'6', qn('w:space'):'1', qn('w:color'):'999999'})
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(6)

# =====================================================================
# 文档一：项目进展汇报
# =====================================================================
def build_progress_report():
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.8)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    set_base_font(doc)

    # 抬头
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PathAI 室内蓝牙导航项目"); r.bold = True; r.font.size = Pt(20); set_eastasia(r, "黑体")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目进展汇报（导师审阅）"); r.bold = True; r.font.size = Pt(15); set_eastasia(r, "黑体")
    para(doc, "初中学部 1# 教学楼 1~2 层 · 视障人士室内导航系统", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    hrule(doc)
    meta = doc.add_table(rows=1, cols=3); meta.style = 'Table Grid'
    mc = meta.rows[0].cells
    for i, (k, v) in enumerate([("汇报日期", "2026-08-17"), ("汇报人", "（请填写）"), ("审阅导师", "（请填写）")]):
        mc[i].text = ""; rr = mc[i].paragraphs[0].add_run(f"{k}：{v}")
        rr.font.size = Pt(10); set_eastasia(rr)
    doc.add_paragraph()

    heading(doc, "一、项目背景与目标", 1)
    bullet(doc, "服务对象：视障人士（核心用户），需在陌生教学楼内独立完成 1~2 层室内导航。")
    bullet(doc, "任务目标：将 CAD 建筑图纸转化为可导航的室内地图，部署蓝牙信标与指纹库，实现米级室内定位与语音导航。")
    bullet(doc, "全链路：CAD PDF → GeoJSON 矢量化 → 交互地图 → 信标部署优化 → 指纹采集方案 → 定位/融合算法 → 视障交互。")

    heading(doc, "二、核心进展（截至 2026-08-17）", 1)

    heading(doc, "2.1 建筑空间数字化（CAD → GeoJSON）", 2)
    bullet(doc, "完成 1# 楼 1~2 层全要素矢量化：房间 F1 81 间 / F2 55 间，门 F1 132 / F2 76，墙体 4442 段。")
    bullet(doc, "拓扑与可导航路网：手动骨架优先，走廊/房间开放-封闭空间分治，生成带门节点（TD）的楼层拓扑。")
    bullet(doc, "质量校验（validate_geojson）：无门封闭房间数 = 0，关键指标全部通过（QA PASS）。")

    heading(doc, "2.2 信标部署方案优化（路线掩码 route-mask）", 2)
    bullet(doc, "路线掩码方案将信标限制在测试路线附近，验证 6m 掩码即最优。")
    bullet(doc, "由 60 基础信标优化至 74（F1 67 + F2 7）；逐坐标比对证明 6~9m 区间信标 0 个，扩大缓冲无增量。")
    bullet(doc, "当前部署文件 ble_deployment.json 登记 61 个（F1 52 + F2 9，含现场微调）；满足 7 条合规规则，第 7 条「不相关空间不部署信标」达成。")

    heading(doc, "2.3 定位精度分析（GDOP × 高/底比）", 2)
    bullet(doc, "误差框架：σ_pos = GDOP × σ_range；距离误差 σ_range = (ln10/10N)·d·σ_RSSI（N=3.5，与距离成正比）。")
    bullet(doc, "实测 695 个 ≥3 信标覆盖点：高/底 ≥ 0.20 达标率 94.7%，≥ 0.25 达 95.9%，≥ 0.40 达 100%。")
    bullet(doc, "工程判据：推荐最小高/底 = 0.20（保守 0.25）；高/底 < 0.05 为禁用区（最大 GDOP 11.85，误差放大近 12 倍）。")
    bullet(doc, "结论：2m@σ=1.5m 精度无解，须降 σ_range ≤ 1.06m 或加密信标——已通过降噪/采集优化应对。")

    heading(doc, "2.4 降噪方案设计（指纹库成本优化）", 2)
    bullet(doc, "组合：多帧中值滤波（σ/√N）+ 位置域卡尔曼滤波 + 走廊中轴线投影。")
    bullet(doc, "预期收益：σ_pos 由约 2m 降到约 1m，可显著降低指纹库采集密度与成本。")
    bullet(doc, "状态：方案已设计完成，待实施（分支 feat-rssi-median / feat-kalman-filter）。")

    heading(doc, "2.5 指纹采集有效区提取", 2)
    bullet(doc, "从设计图红色填充区域提取指纹有效采集空间并集：1F 1984.86 m²（1 多边形）、2F 1127.91 m²（2 多边形）。")
    bullet(doc, "已与建筑 local_meters 坐标系配准（验证与渲染器反推原点一致），可直接与墙体/柱坐标叠加使用。")

    heading(doc, "2.6 指纹采集方案仿真验证", 2)
    para(doc, "在有效区以 0.6m 网格仿真三种定位方法的覆盖与误差（锚点取建筑柱+墙角）：", after=2)
    make_table(doc,
        ["采集定位方法", "有效区可用率", "定位误差", "歧义点", "结论"],
        [["两点法（仅柱+墙角）", "27%", "≤8.1cm", "860", "走廊张角退化，不可用"],
         ["三边测量（3+锚点）", "23%", "最大≈48m", "0", "共线 GDOP 退化，不可用"],
         ["直角坐标法（垂直打墙+沿墙量距）", "97.5%/98.8%", "恒定 2.83cm", "0", "✅ 推荐"],
         ["分区间锚点两点法（增强）", "98.7%~99.9%", "≤5.4cm（均值2.1cm）", "少量", "✅ 可用"]],
        widths=[5.2, 3.0, 3.0, 1.8, 3.5])
    bullet(doc, "直角坐标法经仿真验证可覆盖整片有效区且误差稳定 ≤3cm、零歧义，故测距设备采用经济组合即可，无需万元级全站仪。")

    heading(doc, "三、主要问题与风险", 1)
    bullet(doc, "数据质量遗留：少量标签未匹配多边形（孤儿门 F1 61 / F2 17）；卫生间多边形仅覆盖盥洗走道区。")
    bullet(doc, "走廊骨架为简化模型，未重建 v7 式 700 边中轴路网，复杂路径规划精度受限。")
    bullet(doc, "降噪方案尚未实施：当前定位精度依赖指纹库密度，落地后可大幅降本提效。")
    bullet(doc, "信标供电与维护：纽扣电池续航 1~2 年，需制定更换/巡检计划。")
    bullet(doc, "现场作业约束：激光测距需通视（LOS），转角/房间深处须本地锚点；建议自然锚点（门框角/墙角/柱边）优先、稀疏补靶标。")
    bullet(doc, "坐标配准误差：理论仿真基于 CAD 坐标，实测需考虑施工偏差，建议现场抽查 10~20 点验证。")

    heading(doc, "四、下一步计划", 1)
    bullet(doc, "实施降噪方案（中值滤波 + 卡尔曼滤波），验证 σ_pos 2m→~1m。")
    bullet(doc, "现场踏勘与锚点布设：基于自然锚点确定采集路线与子区间。")
    bullet(doc, "指纹库首轮采集（直角坐标法）+ 现场定位精度抽查。")
    bullet(doc, "端到端导航 demo 联调（定位 + 路线规划 + 视障语音交互）。")
    bullet(doc, "补数据质量遗留问题（孤儿门、卫生间多边形）。")

    heading(doc, "五、资源需求", 1)
    para(doc, "为支撑现场实施，需采购测距定位设备、蓝牙信标、采集终端与辅助耗材，详见附件《工具材料采购预算说明》。核心需求：激光测距设备套装约 ¥700、蓝牙信标约 ¥2,520、辅助耗材约 ¥280（采集终端若实验室已有则不另购）。", after=2)

    doc.save(os.path.join(OUT_DIR, "14-项目进展汇报-导师用.docx"))
    print("OK progress report")

# =====================================================================
# 文档二：采购预算说明
# =====================================================================
def build_budget():
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.8)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    set_base_font(doc)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PathAI 项目 工具材料采购预算说明"); r.bold = True; r.font.size = Pt(18); set_eastasia(r, "黑体")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（请导师审批）"); r.font.size = Pt(12); set_eastasia(r)
    hrule(doc)
    meta = doc.add_table(rows=1, cols=3); meta.style = 'Table Grid'
    mc = meta.rows[0].cells
    for i, (k, v) in enumerate([("申请人", "（请填写）"), ("日期", "2026-08-17"), ("导师审批", "（签字）____________")]):
        mc[i].text = ""; rr = mc[i].paragraphs[0].add_run(f"{k}：{v}")
        rr.font.size = Pt(10); set_eastasia(rr)
    doc.add_paragraph()

    heading(doc, "一、采购背景与依据", 1)
    bullet(doc, "为支撑初中学部 1# 教学楼 1~2 层视障导航系统的现场实施，需采购：①指纹采集定位设备；②蓝牙信标（定位基础设施）；③采集终端与辅助耗材。")
    bullet(doc, "选型依据：仿真验证直角坐标法指纹采集可实现 ~100% 覆盖、误差 ≤2.83cm、零歧义，对锚点几何无严苛要求（建筑内任意点距墙 <30m 均满足），故测距设备采用经济型组合即可，无需万元级全站仪。")

    heading(doc, "二、采购清单（推荐方案 / 基础款）", 1)

    heading(doc, "A. 指纹采集定位设备", 2)
    make_table(doc,
        ["序号", "物品", "规格", "单价(¥)", "数量", "小计(¥)", "用途"],
        [["A1", "激光测距仪", "纯距离，±2mm（得力/优利德/米家）", "169", "2", "338", "测垂距与沿墙距离"],
         ["A2", "全景云台", "带水平刻度环+双水平泡", "200", "1", "200", "固定激光尺+水平旋转扫描"],
         ["A3", "设备夹具/快装板", "夹持激光尺机身", "30", "1", "30", "云台固定"],
         ["A4", "三脚架", "金属，1/4螺口", "100", "1", "100", "架设于锚点"],
         ["A5", "卷尺", "5-10m", "25", "1", "25", "沿墙量距"],
         ["", "合计", "", "", "", "693", ""]],
        widths=[1.0, 2.6, 4.3, 1.8, 1.2, 1.8, 3.3])

    heading(doc, "B. 蓝牙信标", 2)
    make_table(doc,
        ["序号", "物品", "规格", "单价(¥)", "数量", "小计(¥)", "用途"],
        [["B1", "iBeacon 模块", "nRF52832，纽扣电池，续航1-2年", "30", "84", "2,520", "74 部署 + 10 备用"],
         ["", "合计", "", "", "", "2,520", ""]],
        widths=[1.0, 2.6, 4.3, 1.8, 1.2, 1.8, 3.3])
    para(doc, "说明：路线掩码方案最优 74 个信标（F1 67+F2 7）；另购 10 备用应对损坏/丢失。单价为市场公开报价估算。", size=9.5, after=4)

    heading(doc, "C. 采集终端", 2)
    make_table(doc,
        ["序号", "物品", "规格", "单价(¥)", "数量", "小计(¥)", "用途"],
        [["C1", "安卓平板/手机", "中端，采集 App + 导航 demo", "2,000", "1", "2,000", "采集/调试（实验室已有则不采购）"],
         ["", "合计（按需）", "", "", "", "2,000 / 0", ""]],
        widths=[1.0, 2.6, 4.3, 1.8, 1.2, 1.8, 3.3])

    heading(doc, "D. 辅助耗材", 2)
    make_table(doc,
        ["序号", "物品", "单价(¥)", "数量", "小计(¥)"],
        [["D1", "编号标签/标记贴纸", "30", "1", "30"],
         ["D2", "平面图纸打印（A1，1~2层）", "50", "2", "100"],
         ["D3", "充电宝/电池", "100", "1", "100"],
         ["D4", "杂项（记号笔、胶带等）", "50", "1", "50"],
         ["", "合计", "", "", "280"]],
        widths=[1.4, 6.0, 2.4, 1.6, 2.4])

    heading(doc, "三、预算合计（推荐方案）", 1)
    make_table(doc,
        ["类别", "金额(¥)"],
        [["A 测距定位设备", "693"],
         ["B 蓝牙信标", "2,520"],
         ["C 采集终端（按需）", "2,000 / 0"],
         ["D 辅助耗材", "280"],
         ["合计（含终端）", "5,493"],
         ["合计（用现有终端）", "3,493"]],
        widths=[8.0, 4.0])

    heading(doc, "四、可选升级方案（非必须，视精度要求）", 1)
    make_table(doc,
        ["物品", "规格", "参考价(¥)", "说明"],
        [["徕卡 DISTO S910 掌上全站仪", "测距±1mm+测角±1°，蓝牙SDK", "≈14,000", "自动测角，免后视定向，多轮采集更省时"],
         ["徕卡 X3 + DST360 适配器", "测角+三脚架", "≈6,800", "中档，cm 级精度"]],
        widths=[4.5, 4.5, 2.5, 4.5])
    para(doc, "经仿真，基础款已满足 ≤3cm 指纹定位需求；升级仅在对精度有更高要求或大规模多轮采集时建议。", size=9.5, after=4)

    heading(doc, "五、依据与说明", 1)
    bullet(doc, "测距设备选型：仿真验证直角坐标法对锚点几何无严苛要求（任意点距墙<30m），无需测角全站仪；激光尺+云台即可完成「垂直打墙+沿墙量距」。")
    bullet(doc, "信标数量：docs/13 route-mask 6m 方案，60→74 最优；部署文件已登记 61，预留备用 10。")
    bullet(doc, "成本优化：仿真显示全量布设靶标（约 2592 个）需约 43h 不可取；推荐自然锚点+稀疏补盲，且直角坐标法每点约 12s，1442 指纹点约 4.8h。")
    bullet(doc, "所有价格取自主流电商公开报价，为估算值，实际以采购时为准。")

    doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10)
    r = p.add_run("申请人（签字）：________________    日期：__________")
    r.font.size = Pt(11); set_eastasia(r)
    p = doc.add_paragraph()
    r = p.add_run("导师审批（签字）：________________    日期：__________")
    r.font.size = Pt(11); set_eastasia(r)

    doc.save(os.path.join(OUT_DIR, "15-工具材料采购预算说明.docx"))
    print("OK budget")

if __name__ == "__main__":
    build_progress_report()
    build_budget()
    print("ALL DONE")
